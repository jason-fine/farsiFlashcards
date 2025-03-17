import os
from google_images_search import GoogleImagesSearch
from deep_translator import GoogleTranslator
from docx import Document
from urllib.request import urlretrieve
from PIL import Image
import apikey

# API credentials
api_key = apikey.api_key
cse_id = apikey.cse_id

gis = GoogleImagesSearch(api_key, cse_id)
translator = GoogleTranslator(source='en', target='fa')

def search_image(query, max_attempts=5):
    """Search for an image, download it, convert it to JPEG, and return its filename."""
    try:
        gis.search({'q': query, 'num': max_attempts})
        results = gis.results()
        if not results:
            print(f"No images found for {query}")
            return None

        for i, result in enumerate(results):
            image_url = result.url
            image_name = f"{query}_{i}.jpg"  # Store as JPG
            temp_image_name = f"{query}_{i}"  # Temporary file name (might be WebP)

            # Download the image
            urlretrieve(image_url, temp_image_name)
            print(f"Downloaded image for {query}: {temp_image_name}")

            # Convert image to valid format
            valid_image = convert_to_jpg(temp_image_name, image_name)
            if valid_image:
                return image_name  # Return first valid image

            print(f"Invalid image file for {query}, trying next result...")
            os.remove(temp_image_name)  # Remove invalid image

        print(f"No valid images found for {query}")
        return None
    except Exception as e:
        print(f"Error downloading image for {query}: {e}")
        return None

def convert_to_jpg(input_path, output_path):
    """Convert an image to a valid JPEG format."""
    try:
        with Image.open(input_path) as img:
            img = img.convert("RGB")  # Ensure compatibility
            img.save(output_path, "JPEG")
        os.remove(input_path)  # Remove original file after conversion
        return output_path
    except Exception as e:
        print(f"Error converting {input_path}: {e}")
        return None

def create_document(words):
    """Create a Word document with words, images, and translations."""
    doc = Document()
    doc.add_heading('Word List with Images and Translations', 0)

    for word in words:
        print(f"Processing: {word}")

        # Add the word to the document
        doc.add_heading(word, level=1)

        # Search for a valid image
        image_name = search_image(word)
        if image_name:
            doc.add_picture(image_name)  # Add the image to the document
            os.remove(image_name)  # Remove the image after adding
        else:
            doc.add_paragraph("[No valid image found]")  # Placeholder text

        # Translate the word and add the translation
        translation = translator.translate(word)
        doc.add_paragraph(f'Farsi Translation: {translation}')

        # Add a line break between words
        doc.add_paragraph('')

    # Save the document
    doc.save('Word_List_with_Images_and_Translations.docx')
    print("Document created successfully!")

# Example usage
words = ['apple', 'book', 'house']
create_document(words)
